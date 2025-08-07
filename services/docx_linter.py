"""
DocX Jinja Linter Service

This module provides comprehensive linting capabilities for Jinja2 templates embedded in .docx files.
It extracts template content from Word documents and validates Jinja syntax, tag matching, and structure.
"""

import re
import time
from typing import List, Dict, Any, Optional, Tuple
from docx import Document
from jinja2 import Environment, TemplateSyntaxError, select_autoescape
from jinja2.nodes import Node
from jinja2.exceptions import TemplateError
import logging

from models.schemas import (
    LintResult, LintError, LintWarning, LintSummary, LintOptions,
    LintErrorType, LintWarningType, DocxLinterException, 
    InvalidFileFormatException, TemplateSyntaxException, DocumentExtractionException
)

logger = logging.getLogger(__name__)


class DocxJinjaLinterService:
    """
    Service for linting Jinja2 templates embedded in .docx files.
    
    Provides comprehensive validation including:
    - Syntax checking using Jinja2 parser
    - Tag matching (if/endif, for/endfor, etc.)
    - Nested structure validation
    - Undefined variable detection
    - Template quality analysis
    """
    
    def __init__(self):
        """Initialize the linter service with a Jinja2 environment."""
        self.env = Environment(
            autoescape=select_autoescape(['html', 'xml']),
            trim_blocks=True,
            lstrip_blocks=True
        )
        
        # Jinja tag patterns for matching analysis (including docxtpl extensions)
        self.tag_patterns = {
            'block_start': re.compile(r'{%\s*(p|tr|tc|r)?\s*(\w+)(?:\s+[^%]*)?%}'),  # Include p, tr, tc, r prefixes
            'block_end': re.compile(r'{%\s*(p|tr|tc|r)?\s*end(\w+)\s*%}'),  # Include p, tr, tc, r prefixes
            'variable': re.compile(r'{{[^}]*}}'),
            'richtext_variable': re.compile(r'{{r\s+[^}]*}}'),  # RichText variables
            'comment': re.compile(r'{#[^#]*#}'),
            'docxtpl_comment': re.compile(r'{#(p|tr|tc|r)\s+[^#]*#}'),  # DocXTPL paragraph/row/cell comments
            'full_tag': re.compile(r'{[%{#][^}%#]*[%}#]}')
        }
        
        # Tags that require matching end tags (excluding 'set' which doesn't need endset)
        self.paired_tags = {
            'if', 'for', 'with', 'block', 'macro', 'call', 
            'filter', 'trans', 'pluralize', 'raw', 'autoescape'
        }
        
        # Tags that are self-contained (including docxtpl extensions)
        self.standalone_tags = {
            'else', 'elif', 'endif', 'endfor', 'endwith', 'endblock', 
            'endmacro', 'endcall', 'endfilter', 'endtrans', 'endpluralize', 
            'endraw', 'endautoescape', 'include', 'import', 'from', 'extends', 
            'break', 'continue', 'set',  # 'set' is standalone, no endset needed
            # DocXTPL-specific tags
            'cellbg', 'colspan', 'hm', 'vm'  # DocXTPL special tags
        }

    async def lint_docx_file(
        self, 
        file_content: bytes, 
        filename: str,
        options: LintOptions = None
    ) -> LintResult:
        """
        Main linting method that processes a .docx file and returns comprehensive results.
        
        Args:
            file_content: Raw bytes of the .docx file
            filename: Original filename for error reporting
            options: Linting configuration options
            
        Returns:
            LintResult containing all errors, warnings, and summary information
        """
        start_time = time.time()
        
        if options is None:
            options = LintOptions()
        
        errors = []
        warnings = []
        
        try:
            # Stage 1: Extract template content
            logger.info(f"Extracting template content from {filename}")
            template_content = self._extract_template_content(file_content, filename)
            
            if not template_content or not template_content.strip():
                errors.append(LintError(
                    error_type=LintErrorType.DOCUMENT_ERROR,
                    message="No template content found in document",
                    suggestion="Ensure the document contains Jinja2 template syntax"
                ))
                
            # Stage 2: Basic statistics
            lines = template_content.split('\n')
            lines_count = len(lines)
            jinja_tags = self._count_jinja_tags(template_content)
            
            # Stage 3: Syntax validation
            if options.check_undefined_vars or template_content.strip():
                syntax_errors = self._validate_jinja_syntax(template_content, options)
                errors.extend(syntax_errors)
            
            # Stage 4: Tag matching validation
            if options.check_tag_matching:
                tag_errors = self._check_tag_matching(template_content, options)
                errors.extend(tag_errors)
            
            # Stage 5: Structure validation
            if options.check_nested_structure:
                structure_errors = self._check_template_structure(template_content, options)
                errors.extend(structure_errors)
            
            # Stage 6: Quality checks and warnings
            quality_warnings = self._check_template_quality(template_content, lines, options)
            warnings.extend(quality_warnings)
            
            # Stage 7: Calculate completeness score
            completeness_score = self._calculate_completeness_score(
                template_content, len(errors), len(warnings), jinja_tags
            )
            
            processing_time = (time.time() - start_time) * 1000
            
            # Create summary
            summary = LintSummary(
                total_errors=len(errors),
                total_warnings=len(warnings),
                template_size=len(template_content),
                lines_count=lines_count,
                jinja_tags_count=jinja_tags,
                completeness_score=completeness_score,
                processing_time_ms=round(processing_time, 2)
            )
            
            # Create preview (first 500 chars)
            preview = template_content[:500] if len(template_content) > 500 else template_content
            if len(template_content) > 500:
                preview += "..."
            
            success = len(errors) == 0 and (not options.fail_on_warnings or len(warnings) == 0)
            
            result = LintResult(
                success=success,
                errors=errors,
                warnings=warnings,
                summary=summary,
                template_content=template_content if options.verbose else None,
                template_preview=preview
            )
            
            logger.info(f"Linting completed: {len(errors)} errors, {len(warnings)} warnings")
            return result
            
        except Exception as e:
            logger.error(f"Linting failed for {filename}: {str(e)}")
            
            # Create error result
            processing_time = (time.time() - start_time) * 1000
            
            if isinstance(e, DocxLinterException):
                errors.append(LintError(
                    error_type=LintErrorType.DOCUMENT_ERROR,
                    message=str(e),
                    suggestion="Check document format and content"
                ))
            else:
                errors.append(LintError(
                    error_type=LintErrorType.DOCUMENT_ERROR,
                    message=f"Unexpected error during linting: {str(e)}",
                    suggestion="Please check the document format and try again"
                ))
            
            summary = LintSummary(
                total_errors=len(errors),
                total_warnings=0,
                template_size=0,
                lines_count=0,
                jinja_tags_count=0,
                processing_time_ms=round(processing_time, 2)
            )
            
            return LintResult(
                success=False,
                errors=errors,
                warnings=[],
                summary=summary,
                template_preview=None
            )

    def _extract_template_content(self, file_content: bytes, filename: str) -> str:
        """
        Extract Jinja template content from .docx file.
        
        Args:
            file_content: Raw bytes of the .docx file
            filename: Original filename for error reporting
            
        Returns:
            Extracted template content as string
            
        Raises:
            DocumentExtractionException: If extraction fails
        """
        try:
            import tempfile
            import os
            
            # Create temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
                temp_file.write(file_content)
                temp_file_path = temp_file.name
            
            try:
                # Load document
                doc = Document(temp_file_path)
                content_parts = []
                
                # Extract text from paragraphs
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        content_parts.append(paragraph.text)
                
                # Extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                content_parts.append(cell.text)
                
                # Extract text from headers and footers
                for section in doc.sections:
                    # Headers
                    if section.header:
                        for paragraph in section.header.paragraphs:
                            if paragraph.text.strip():
                                content_parts.append(paragraph.text)
                    
                    # Footers
                    if section.footer:
                        for paragraph in section.footer.paragraphs:
                            if paragraph.text.strip():
                                content_parts.append(paragraph.text)
                
                # Combine all content
                template_content = '\n'.join(content_parts)
                
                logger.info(f"Extracted {len(template_content)} characters from {filename}")
                return template_content
                
            finally:
                # Clean up temporary file
                if os.path.exists(temp_file_path):
                    os.unlink(temp_file_path)
                    
        except Exception as e:
            raise DocumentExtractionException(f"Failed to extract content from {filename}: {str(e)}")

    def _validate_jinja_syntax(self, content: str, options: LintOptions) -> List[LintError]:
        """
        Validate Jinja2 syntax using the Jinja2 parser.
        
        Args:
            content: Template content to validate
            options: Linting options
            
        Returns:
            List of syntax errors found
        """
        errors = []
        
        try:
            # Try to parse the template
            self.env.parse(content)
            logger.debug("Jinja2 syntax validation passed")
            
        except TemplateSyntaxError as e:
            error = LintError(
                line_number=getattr(e, 'lineno', None),
                column=getattr(e, 'colno', None),
                error_type=LintErrorType.SYNTAX_ERROR,
                message=f"Jinja2 syntax error: {str(e)}",
                context=self._get_line_context(content, getattr(e, 'lineno', None)),
                suggestion="Check Jinja2 syntax documentation for correct tag format"
            )
            errors.append(error)
            logger.warning(f"Syntax error at line {getattr(e, 'lineno', 'unknown')}: {str(e)}")
            
        except TemplateError as e:
            error = LintError(
                error_type=LintErrorType.SYNTAX_ERROR,
                message=f"Template error: {str(e)}",
                suggestion="Check template for syntax issues"
            )
            errors.append(error)
            logger.warning(f"Template error: {str(e)}")
            
        return errors

    def _check_tag_matching(self, content: str, options: LintOptions) -> List[LintError]:
        """
        Check for matching Jinja tag pairs (if/endif, for/endfor, etc.).
        
        Args:
            content: Template content to check
            options: Linting options
            
        Returns:
            List of tag matching errors
        """
        errors = []
        lines = content.split('\n')
        tag_stack = []  # Stack to track opening tags
        
        for line_num, line in enumerate(lines, 1):
            # Find all Jinja tags in the line
            block_starts = self.tag_patterns['block_start'].finditer(line)
            block_ends = self.tag_patterns['block_end'].finditer(line)
            
            # Process opening tags
            for match in block_starts:
                prefix = match.group(1) or ''  # p, tr, tc, r prefix (may be empty)
                tag_name = match.group(2).lower() if match.group(2) else match.group(1).lower()
                full_match = match.group(0)
                
                if tag_name in self.paired_tags:
                    tag_stack.append({
                        'tag': tag_name,
                        'prefix': prefix,
                        'line': line_num,
                        'content': full_match.strip(),
                        'position': match.start()
                    })
                elif tag_name not in self.standalone_tags:
                    # Unknown tag (but don't flag docxtpl prefixed tags as unknown)
                    if not prefix or prefix not in ['p', 'tr', 'tc', 'r']:
                        errors.append(LintError(
                        line_number=line_num,
                        column=match.start(),
                        error_type=LintErrorType.SYNTAX_ERROR,
                        message=f"Unknown Jinja tag: {tag_name}",
                        context=line.strip(),
                        tag_name=tag_name,
                        suggestion="Check if tag name is spelled correctly"
                    ))
            
            # Process closing tags
            for match in block_ends:
                end_prefix = match.group(1) or ''  # p, tr, tc, r prefix (may be empty)
                end_tag_name = match.group(2).lower() if match.group(2) else match.group(1).lower()
                full_match = match.group(0)
                
                if not tag_stack:
                    # Closing tag without opening
                    errors.append(LintError(
                        line_number=line_num,
                        column=match.start(),
                        error_type=LintErrorType.MISMATCHED_TAG,
                        message=f"Closing tag '{end_tag_name}' without matching opening tag",
                        context=line.strip(),
                        tag_name=end_tag_name,
                        suggestion=f"Add opening {{% {end_tag_name} %}} tag before this line"
                    ))
                    continue
                
                # Check if closing tag matches the most recent opening tag
                expected_tag = tag_stack[-1]['tag']
                expected_prefix = tag_stack[-1].get('prefix', '')
                if end_tag_name == expected_tag and end_prefix == expected_prefix:
                    tag_stack.pop()  # Correct match
                else:
                    # Mismatched tags
                    opening_info = tag_stack[-1]
                    expected_full = f"{expected_prefix}end{expected_tag}" if expected_prefix else f"end{expected_tag}"
                    found_full = f"{end_prefix}end{end_tag_name}" if end_prefix else f"end{end_tag_name}"
                    errors.append(LintError(
                        line_number=line_num,
                        column=match.start(),
                        error_type=LintErrorType.MISMATCHED_TAG,
                        message=f"Expected '{expected_full}' but found '{found_full}'",
                        context=line.strip(),
                        tag_name=end_tag_name,
                        suggestion=f"Change to {{% {expected_full} %}} or check tag nesting (opened at line {opening_info['line']})"
                    ))
        
        # Check for unclosed tags
        for unclosed_tag in tag_stack:
            tag_prefix = unclosed_tag.get('prefix', '')
            tag_name = unclosed_tag['tag']
            full_tag = f"{tag_prefix}{tag_name}" if tag_prefix else tag_name
            close_tag = f"{tag_prefix}end{tag_name}" if tag_prefix else f"end{tag_name}"
            
            errors.append(LintError(
                line_number=unclosed_tag['line'],
                error_type=LintErrorType.UNCLOSED_TAG,
                message=f"Unclosed '{full_tag}' tag",
                context=unclosed_tag['content'],
                tag_name=tag_name,
                suggestion=f"Add {{% {close_tag} %}} tag to close this block"
            ))
        
        logger.debug(f"Tag matching check found {len(errors)} errors")
        return errors

    def _check_template_structure(self, content: str, options: LintOptions) -> List[LintError]:
        """
        Check nested structure and tag ordering.
        
        Args:
            content: Template content to check
            options: Linting options
            
        Returns:
            List of structure errors
        """
        errors = []
        lines = content.split('\n')
        
        # Track nesting depth and context
        nesting_stack = []
        max_depth = 0
        current_depth = 0
        
        for line_num, line in enumerate(lines, 1):
            # Find block tags
            block_starts = self.tag_patterns['block_start'].finditer(line)
            block_ends = self.tag_patterns['block_end'].finditer(line)
            
            # Process opening tags
            for match in block_starts:
                prefix = match.group(1) or ''  # p, tr, tc, r prefix (may be empty)
                tag_name = match.group(2).lower() if match.group(2) else match.group(1).lower()
                
                if tag_name in self.paired_tags:
                    current_depth += 1
                    max_depth = max(max_depth, current_depth)
                    
                    nesting_stack.append({
                        'tag': tag_name,
                        'line': line_num,
                        'depth': current_depth
                    })
                    
                    # Check for excessive nesting
                    if current_depth > 10:
                        errors.append(LintError(
                            line_number=line_num,
                            error_type=LintErrorType.NESTED_ERROR,
                            message=f"Excessive nesting depth ({current_depth})",
                            context=line.strip(),
                            tag_name=tag_name,
                            suggestion="Consider breaking complex logic into smaller templates"
                        ))
            
            # Process closing tags
            for match in block_ends:
                end_prefix = match.group(1) or ''  # p, tr, tc, r prefix (may be empty)
                end_tag_name = match.group(2).lower() if match.group(2) else match.group(1).lower()
                
                if nesting_stack and end_tag_name == nesting_stack[-1]['tag']:
                    nesting_stack.pop()
                    current_depth -= 1
        
        logger.debug(f"Structure check found {len(errors)} errors, max depth: {max_depth}")
        return errors

    def _check_template_quality(self, content: str, lines: List[str], options: LintOptions) -> List[LintWarning]:
        """
        Perform quality checks and generate warnings.
        
        Args:
            content: Template content
            lines: List of template lines
            options: Linting options
            
        Returns:
            List of quality warnings
        """
        warnings = []
        
        # Check line length
        for line_num, line in enumerate(lines, 1):
            if len(line) > options.max_line_length:
                warnings.append(LintWarning(
                    line_number=line_num,
                    warning_type=LintWarningType.LONG_LINE,
                    message=f"Line too long ({len(line)} > {options.max_line_length} characters)",
                    context=line[:100] + "..." if len(line) > 100 else line,
                    suggestion="Consider breaking long lines for better readability"
                ))
        
        # Check for complex expressions
        variable_matches = self.tag_patterns['variable'].finditer(content)
        for match in variable_matches:
            var_content = match.group(0)
            if len(var_content) > 50:  # Arbitrary threshold for complexity
                line_num = content[:match.start()].count('\n') + 1
                warnings.append(LintWarning(
                    line_number=line_num,
                    warning_type=LintWarningType.COMPLEX_EXPRESSION,
                    message="Complex variable expression detected",
                    context=var_content,
                    suggestion="Consider simplifying expression or using intermediate variables"
                ))
        
        # Check for suspicious patterns
        suspicious_patterns = [
            (r'{{{[^}]*}}}', "Triple braces detected"),
            (r'{%\s*{[^}]*}\s*%}', "Mixed tag syntax"),
            (r'{{[^}]*{{[^}]*}}', "Nested double braces in variable"),
            (r'{\s+{[^}]*}\s+}', "Spaces between braces")
        ]
        
        for pattern, description in suspicious_patterns:
            matches = re.finditer(pattern, content)
            for match in matches:
                line_num = content[:match.start()].count('\n') + 1
                warnings.append(LintWarning(
                    line_number=line_num,
                    warning_type=LintWarningType.SUSPICIOUS_SYNTAX,
                    message=f"Suspicious syntax: {description}",
                    context=match.group(0),
                    suggestion="Review syntax for correctness"
                ))
        
        logger.debug(f"Quality check generated {len(warnings)} warnings")
        return warnings

    def _count_jinja_tags(self, content: str) -> int:
        """Count the number of Jinja tags in the template."""
        return len(self.tag_patterns['full_tag'].findall(content))

    def _get_line_context(self, content: str, line_number: Optional[int], context_lines: int = 2) -> Optional[str]:
        """
        Get context lines around a specific line number.
        
        Args:
            content: Template content
            line_number: Target line number
            context_lines: Number of lines before and after to include
            
        Returns:
            Context string or None if line_number is None
        """
        if line_number is None:
            return None
        
        lines = content.split('\n')
        start = max(0, line_number - context_lines - 1)
        end = min(len(lines), line_number + context_lines)
        
        context = []
        for i in range(start, end):
            marker = " -> " if i == line_number - 1 else "    "
            context.append(f"{marker}{i+1}: {lines[i]}")
        
        return '\n'.join(context)

    def _calculate_completeness_score(self, content: str, errors: int, warnings: int, tags: int) -> float:
        """
        Calculate a completeness/quality score for the template.
        
        Args:
            content: Template content
            errors: Number of errors
            warnings: Number of warnings
            tags: Number of Jinja tags
            
        Returns:
            Score from 0-100
        """
        if not content.strip():
            return 0.0
        
        # Base score
        score = 100.0
        
        # Subtract for errors (more impactful)
        score -= (errors * 15)
        
        # Subtract for warnings (less impactful)
        score -= (warnings * 5)
        
        # Bonus for having Jinja tags (shows it's actually a template)
        if tags > 0:
            score += min(tags * 2, 10)  # Max 10 bonus points
        
        # Ensure score is within bounds
        return max(0.0, min(100.0, score))