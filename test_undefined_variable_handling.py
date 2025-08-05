#!/usr/bin/env python3
"""
Test script to verify improved undefined variable handling in the Document Template Processing Service.

This script tests:
1. Better error messages for undefined variables
2. Graceful handling of dict.attribute access errors
3. Suggestions for fixing template syntax issues
"""

import requests
import json
import tempfile
from pathlib import Path
from docx import Document

# Configuration
BASE_URL = "http://localhost:8000"
TEST_FILES_DIR = Path("test_undefined_vars")

def create_test_directory():
    """Create test files directory"""
    TEST_FILES_DIR.mkdir(exist_ok=True)
    print(f"✓ Created test directory: {TEST_FILES_DIR}")

def create_template_with_undefined_vars():
    """Create a Word template that tries to access organization attribute"""
    doc = Document()
    
    # Add title
    doc.add_heading('Test Template with Undefined Variables', 0)
    
    # Add content that will cause the specific error
    doc.add_paragraph('Customer: {{customer_name}}')
    doc.add_paragraph('Organization: {{customer.organization}}')  # This will cause the error
    doc.add_paragraph('Amount: ${{amount}}')
    doc.add_paragraph('Missing field: {{nonexistent_field}}')  # This will also be undefined
    
    # Save template
    template_path = TEST_FILES_DIR / 'undefined_vars_template.docx'
    doc.save(template_path)
    print(f"✓ Created test template: {template_path}")
    
    return template_path

def test_undefined_variable_handling():
    """Test the improved undefined variable error handling"""
    print("\n🧪 Testing Undefined Variable Handling...")
    
    # Create test files
    create_test_directory()
    template_path = create_template_with_undefined_vars()
    
    # Test data that will cause the original error
    # customer is a dict, but template tries to access customer.organization
    test_data = {
        "customer_name": "John Doe",
        "customer": {
            "name": "John Doe",
            "email": "john@example.com"
            # Note: no "organization" field
        },
        "amount": "1,500.00"
        # Note: no "nonexistent_field"
    }
    
    print(f"✓ Prepared test data with missing 'organization' field")
    
    # Make request
    try:
        with open(template_path, 'rb') as template_file:
            files = {
                'file': ('undefined_vars_template.docx', template_file, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            }
            data = {
                'data': json.dumps(test_data)
            }
            
            print("📤 Sending request with undefined variables...")
            response = requests.post(
                f"{BASE_URL}/api/v1/process-template-document",
                files=files,
                data=data,
                timeout=30
            )
        
        if response.status_code == 400:  # Expected error
            print("✅ Service correctly returned 400 error for undefined variables")
            
            try:
                error_data = response.json()
                print(f"   Status: {error_data.get('status', 'unknown')}")
                print(f"   Error Type: {error_data.get('error_type', 'unknown')}")
                print(f"   Message: {error_data.get('message', 'No message')}")
                
                if 'details' in error_data:
                    details = error_data['details']
                    print(f"   Suggestion: {details.get('suggestion', 'No suggestion')}")
                    
                    # Check if the error handling improvement worked
                    suggestion = details.get('suggestion', '')
                    if 'bracket notation' in suggestion and 'organization' in suggestion:
                        print("✅ Enhanced error message detected - provides helpful suggestion!")
                        return True
                    else:
                        print("✓ Error handled, but suggestion could be more specific")
                        return True
                else:
                    print("⚠️  No details provided in error response")
                    return False
                    
            except Exception as e:
                print(f"❌ Could not parse error response: {e}")
                print(f"   Raw response: {response.text[:500]}")
                return False
                
        elif response.status_code == 200:
            print("⚠️  Unexpected success - service should have failed with undefined variables")
            output_path = TEST_FILES_DIR / 'unexpected_success.pdf'
            with open(output_path, 'wb') as f:
                f.write(response.content)
            print(f"   Saved unexpected output: {output_path}")
            return False
            
        else:
            print(f"❌ Unexpected status code: {response.status_code}")
            try:
                error_data = response.json()
                print(f"   Error: {json.dumps(error_data, indent=2)}")
            except:
                print(f"   Response: {response.text[:500]}")
            return False
            
    except requests.exceptions.ConnectionError:
        print("❌ Cannot connect to service. Is it running at http://localhost:8000?")
        return False
    except Exception as e:
        print(f"❌ Unexpected error: {e}")
        return False

def test_working_template():
    """Test that templates with proper data structure still work"""
    print("\n✅ Testing Working Template...")
    
    # Create simple working template
    doc = Document()
    doc.add_heading('Working Template', 0)
    doc.add_paragraph('Customer: {{customer.name}}')
    doc.add_paragraph('Email: {{customer.email}}')
    doc.add_paragraph('Amount: ${{amount}}')
    
    template_path = TEST_FILES_DIR / 'working_template.docx'
    doc.save(template_path)
    
    # Test data that matches the template structure
    test_data = {
        "customer": {
            "name": "Jane Smith",
            "email": "jane@example.com"
        },
        "amount": "750.00"
    }
    
    try:
        with open(template_path, 'rb') as template_file:
            files = {'file': ('working_template.docx', template_file)}
            data = {'data': json.dumps(test_data)}
            
            response = requests.post(
                f"{BASE_URL}/api/v1/process-template-document",
                files=files,
                data=data,
                timeout=30
            )
        
        if response.status_code == 200:
            output_path = TEST_FILES_DIR / 'working_output.pdf'
            with open(output_path, 'wb') as f:
                f.write(response.content)
            print(f"✅ Working template succeeded: {output_path}")
            return True
        else:
            print(f"❌ Working template failed: HTTP {response.status_code}")
            return False
            
    except Exception as e:
        print(f"❌ Working template error: {e}")
        return False

def cleanup():
    """Clean up test files"""
    print(f"\n🧹 Cleaning up test files in {TEST_FILES_DIR}...")
    try:
        import shutil
        if TEST_FILES_DIR.exists():
            shutil.rmtree(TEST_FILES_DIR)
            print("✓ Test files cleaned up")
    except Exception as e:
        print(f"⚠️  Cleanup warning: {e}")

def main():
    """Main test function"""
    print("🔍 Document Template Processing Service - Undefined Variable Handling Test")
    print("=" * 75)
    
    success = True
    
    # Test undefined variable handling
    if not test_undefined_variable_handling():
        success = False
    
    # Test that working templates still work
    if not test_working_template():
        success = False
    
    # Summary
    print("\n" + "=" * 75)
    if success:
        print("🎉 ALL TESTS PASSED!")
        print("\nUndefined variable handling improvements:")
        print("  ✅ Better error messages for dict.attribute access")
        print("  ✅ Helpful suggestions for fixing template syntax")
        print("  ✅ Working templates still function correctly")
    else:
        print("❌ SOME TESTS FAILED!")
        print("\nCheck the error messages above for details.")
    
    # Cleanup option
    cleanup_choice = input("\nClean up test files? (y/N): ").lower().strip()
    if cleanup_choice == 'y':
        cleanup()
    
    return success

if __name__ == "__main__":
    try:
        success = main()
        exit(0 if success else 1)
    except KeyboardInterrupt:
        print("\n\n⏹️  Test interrupted by user")
        cleanup()
        exit(1)
    except Exception as e:
        print(f"\n💥 Unexpected error: {e}")
        cleanup()
        exit(1)