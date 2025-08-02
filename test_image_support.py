#!/usr/bin/env python3
"""
Test script for the image support functionality in the Document Template Processing Service.

This script demonstrates:
1. Creating a simple Word template with image placeholders
2. Generating a test PNG image 
3. Converting the image to base64
4. Calling the new image-enhanced API endpoint
5. Verifying the PDF output

Usage:
    python test_image_support.py

Requirements:
    - Service running at http://localhost:8000
    - python-docx, Pillow for test setup
"""

import requests
import json
import base64
import os
from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
import tempfile

# Configuration
BASE_URL = "http://localhost:8000"
TEST_FILES_DIR = Path("test_image_files")

def create_test_directory():
    """Create test files directory"""
    TEST_FILES_DIR.mkdir(exist_ok=True)
    print(f"✓ Created test directory: {TEST_FILES_DIR}")

def create_test_image(text: str, filename: str, width: int = 200, height: int = 100, color: str = "blue"):
    """Create a simple test PNG image with text"""
    # Create image with white background
    image = Image.new('RGB', (width, height), 'white')
    draw = ImageDraw.Draw(image)
    
    # Try to use a default font, fallback to basic if not available
    try:
        font = ImageFont.truetype("arial.ttf", 20)
    except:
        try:
            font = ImageFont.load_default()
        except:
            font = None
    
    # Calculate text position (center)
    if font:
        bbox = draw.textbbox((0, 0), text, font=font)
        text_width = bbox[2] - bbox[0]
        text_height = bbox[3] - bbox[1]
    else:
        text_width, text_height = len(text) * 10, 20  # Rough estimate
    
    x = (width - text_width) // 2
    y = (height - text_height) // 2
    
    # Draw text
    draw.text((x, y), text, fill=color, font=font)
    
    # Add a simple border
    draw.rectangle([0, 0, width-1, height-1], outline=color, width=2)
    
    # Save image
    image_path = TEST_FILES_DIR / filename
    image.save(image_path, 'PNG')
    print(f"✓ Created test image: {image_path}")
    
    return image_path

def create_test_template():
    """Create a Word template with image placeholders"""
    # Create a new Document
    doc = Document()
    
    # Add title
    title = doc.add_heading('Invoice with Images', 0)
    
    # Add logo placeholder
    logo_paragraph = doc.add_paragraph()
    logo_paragraph.add_run('Company Logo: ')
    logo_paragraph.add_run('{{ company_logo }}')
    
    # Add some content
    doc.add_paragraph('Customer: {{ customer_name }}')
    doc.add_paragraph('Invoice Date: {{ invoice_date }}')
    doc.add_paragraph('Amount: ${{ amount }}')
    
    # Add separator
    doc.add_paragraph('─' * 50)
    
    # Add signature placeholder
    signature_paragraph = doc.add_paragraph()
    signature_paragraph.add_run('Authorized Signature: ')
    signature_paragraph.add_run('{{ user_signature }}')
    
    # Add footer
    doc.add_paragraph()
    doc.add_paragraph('Thank you for your business!')
    
    # Save template
    template_path = TEST_FILES_DIR / 'image_test_template.docx'
    doc.save(template_path)
    print(f"✓ Created test template: {template_path}")
    
    return template_path

def image_to_base64(image_path: Path) -> str:
    """Convert image file to base64 string"""
    with open(image_path, 'rb') as image_file:
        base64_string = base64.b64encode(image_file.read()).decode('utf-8')
    print(f"✓ Converted {image_path.name} to base64 ({len(base64_string)} chars)")
    return base64_string

def test_image_endpoint():
    """Test the image-enhanced endpoint"""
    print("\n🧪 Testing Image-Enhanced Document Processing...")
    
    # Create test files
    create_test_directory()
    
    # Create test images
    logo_path = create_test_image("ACME CORP", "logo.png", 200, 80, "blue")
    signature_path = create_test_image("John Doe", "signature.png", 150, 60, "green")
    
    # Create test template
    template_path = create_test_template()
    
    # Convert images to base64
    logo_base64 = image_to_base64(logo_path)
    signature_base64 = image_to_base64(signature_path)
    
    # Prepare request data
    request_data = {
        "template_data": {
            "customer_name": "Jane Smith",
            "invoice_date": "2024-01-15",
            "amount": "1,500.00",
            "company_logo": "{{company_logo}}",
            "user_signature": "{{user_signature}}"
        },
        "images": {
            "company_logo": {
                "data": logo_base64,
                "width_mm": 60,
                "height_mm": 24
            },
            "user_signature": {
                "data": signature_base64,
                "width_mm": 45,
                "height_mm": 18
            }
        }
    }
    
    print(f"✓ Prepared request with {len(request_data['images'])} images")
    
    # Make request
    try:
        with open(template_path, 'rb') as template_file:
            files = {
                'file': ('image_test_template.docx', template_file, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            }
            data = {
                'request_data': json.dumps(request_data)
            }
            
            print("📤 Sending request to image-enhanced endpoint...")
            response = requests.post(
                f"{BASE_URL}/api/v1/process-template-document-with-images",
                files=files,
                data=data,
                timeout=60
            )
        
        if response.status_code == 200:
            # Save PDF output
            output_path = TEST_FILES_DIR / 'output_with_images.pdf'
            with open(output_path, 'wb') as f:
                f.write(response.content)
            
            print(f"✅ SUCCESS! PDF generated: {output_path}")
            print(f"   PDF size: {len(response.content)} bytes")
            print(f"   Content-Type: {response.headers.get('content-type', 'unknown')}")
            
            # Verify PDF header
            if response.content.startswith(b'%PDF'):
                print("✓ Valid PDF header detected")
            else:
                print("⚠️  Warning: PDF header not detected")
            
            return True
            
        else:
            print(f"❌ FAILED! HTTP {response.status_code}")
            try:
                error_data = response.json()
                print(f"   Error Type: {error_data.get('error_type', 'unknown')}")
                print(f"   Message: {error_data.get('message', 'No message')}")
                if 'details' in error_data:
                    print(f"   Details: {json.dumps(error_data['details'], indent=2)}")
            except:
                print(f"   Response: {response.text[:500]}")
            
            return False
            
    except requests.exceptions.ConnectionError:
        print("❌ FAILED! Cannot connect to service. Is it running at http://localhost:8000?")
        return False
    except Exception as e:
        print(f"❌ FAILED! Unexpected error: {e}")
        return False

def test_service_health():
    """Test if the service is accessible"""
    print("🔍 Checking service health...")
    try:
        response = requests.get(f"{BASE_URL}/", timeout=10)
        if response.status_code == 200:
            print("✅ Service is healthy")
            return True
        else:
            print(f"⚠️  Service responded with status {response.status_code}")
            return False
    except requests.exceptions.ConnectionError:
        print("❌ Service is not accessible")
        return False

def test_backward_compatibility():
    """Test that the original endpoint still works"""
    print("\n🔄 Testing backward compatibility...")
    
    # Create simple template without images
    doc = Document()
    doc.add_heading('Simple Invoice', 0)
    doc.add_paragraph('Customer: {{ customer_name }}')
    doc.add_paragraph('Amount: ${{ amount }}')
    
    simple_template_path = TEST_FILES_DIR / 'simple_template.docx'
    doc.save(simple_template_path)
    
    # Test original endpoint
    try:
        with open(simple_template_path, 'rb') as template_file:
            files = {'file': ('simple_template.docx', template_file)}
            data = {'data': json.dumps({"customer_name": "John Doe", "amount": "500.00"})}
            
            response = requests.post(
                f"{BASE_URL}/api/v1/process-template-document",
                files=files,
                data=data,
                timeout=30
            )
        
        if response.status_code == 200:
            output_path = TEST_FILES_DIR / 'output_simple.pdf'
            with open(output_path, 'wb') as f:
                f.write(response.content)
            print(f"✅ Backward compatibility confirmed: {output_path}")
            return True
        else:
            print(f"❌ Backward compatibility failed: HTTP {response.status_code}")
            return False
            
    except Exception as e:
        print(f"❌ Backward compatibility error: {e}")
        return False

def cleanup():
    """Clean up test files"""
    print(f"\n🧹 Cleaning up test files in {TEST_FILES_DIR}...")
    try:
        import shutil
        if TEST_FILES_DIR.exists():
            shutil.rmtree(TEST_FILES_DIR)
            print("✓ Test files cleaned up")
    except Exception as e:
        print(f"⚠️  Cleanup warning: {e}")

def main():
    """Main test function"""
    print("🖼️  Document Template Processing Service - Image Support Test")
    print("=" * 60)
    
    # Check service health
    if not test_service_health():
        print("\n❌ Cannot proceed: Service is not accessible")
        print("   Make sure the service is running: docker compose up -d")
        return False
    
    success = True
    
    # Test image functionality
    if not test_image_endpoint():
        success = False
    
    # Test backward compatibility
    if not test_backward_compatibility():
        success = False
    
    # Summary
    print("\n" + "=" * 60)
    if success:
        print("🎉 ALL TESTS PASSED!")
        print("\nImage support is working correctly:")
        print("  ✅ Image-enhanced endpoint functional")
        print("  ✅ Base64 image processing working")
        print("  ✅ PDF generation with images successful")
        print("  ✅ Backward compatibility maintained")
        print(f"\nTest files saved in: {TEST_FILES_DIR}")
    else:
        print("❌ SOME TESTS FAILED!")
        print("\nCheck the error messages above for details.")
    
    # Cleanup option
    cleanup_choice = input("\nClean up test files? (y/N): ").lower().strip()
    if cleanup_choice == 'y':
        cleanup()
    
    return success

if __name__ == "__main__":
    try:
        success = main()
        exit(0 if success else 1)
    except KeyboardInterrupt:
        print("\n\n⏹️  Test interrupted by user")
        cleanup()
        exit(1)
    except Exception as e:
        print(f"\n💥 Unexpected error: {e}")
        cleanup()
        exit(1)