"""
Unit and integration tests for the DocX Jinja Linter Service.
"""

import pytest
import asyncio
import tempfile
import os
import json
from typing import Dict, Any
from fastapi.testclient import TestClient
from docx import Document

from main import app
from services.docx_linter import DocxJinjaLinterService
from models.schemas import LintOptions, LintErrorType, LintWarningType


class TestDocxJinjaLinterService:
    """Test cases for the core linter service."""
    
    def setup_method(self):
        """Set up test fixtures."""
        self.linter = DocxJinjaLinterService()
    
    def create_test_docx(self, content: str) -> bytes:
        """
        Create a test .docx file with the given content.
        
        Args:
            content: Text content to put in the document
            
        Returns:
            Raw bytes of the .docx file
        """
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
            temp_path = temp_file.name
        
        try:
            doc = Document()
            # Split content into paragraphs and add to document
            paragraphs = content.split('\n')
            for paragraph_text in paragraphs:
                if paragraph_text.strip():  # Skip empty lines
                    doc.add_paragraph(paragraph_text)
            
            doc.save(temp_path)
            
            # Read the file as bytes
            with open(temp_path, 'rb') as f:
                return f.read()
        finally:
            # Clean up temporary file
            if os.path.exists(temp_path):
                os.unlink(temp_path)
    
    @pytest.mark.asyncio
    async def test_valid_template_no_errors(self):
        """Test linting a valid template with no errors."""
        content = """
        Hello {{ name }}!
        
        {% if company %}
        Company: {{ company }}
        {% endif %}
        
        {% for item in items %}
        - {{ item.name }}: {{ item.price }}
        {% endfor %}
        """
        
        docx_bytes = self.create_test_docx(content)
        result = await self.linter.lint_docx_file(docx_bytes, "test.docx")
        
        assert result.success is True
        assert len(result.errors) == 0
        assert result.summary.total_errors == 0
        assert result.summary.jinja_tags_count > 0
        assert result.template_preview is not None
    
    @pytest.mark.asyncio
    async def test_unclosed_if_tag(self):
        """Test detection of unclosed if tag."""
        content = """
        Hello {{ name }}!
        {% if company %}
        Company: {{ company }}
        # Missing {% endif %}
        """
        
        docx_bytes = self.create_test_docx(content)
        result = await self.linter.lint_docx_file(docx_bytes, "test.docx")
        
        assert result.success is False
        assert len(result.errors) > 0
        
        # Check for unclosed tag error
        unclosed_errors = [e for e in result.errors if e.error_type == LintErrorType.UNCLOSED_TAG]
        assert len(unclosed_errors) > 0
        assert unclosed_errors[0].tag_name == "if"
    
    @pytest.mark.asyncio
    async def test_mismatched_tags(self):
        """Test detection of mismatched tags."""
        content = """
        {% if condition %}
        Some content
        {% endfor %}  # Should be {% endif %}
        """
        
        docx_bytes = self.create_test_docx(content)
        result = await self.linter.lint_docx_file(docx_bytes, "test.docx")
        
        assert result.success is False
        assert len(result.errors) > 0
        
        # Check for mismatched tag error
        mismatch_errors = [e for e in result.errors if e.error_type == LintErrorType.MISMATCHED_TAG]
        assert len(mismatch_errors) > 0
    
    @pytest.mark.asyncio 
    async def test_syntax_error(self):
        """Test detection of Jinja syntax errors."""
        content = """
        Hello {{ name }!  # Missing closing brace
        {% if condition %  # Missing closing brace
        Content
        {% endif %}
        """
        
        docx_bytes = self.create_test_docx(content)
        result = await self.linter.lint_docx_file(docx_bytes, "test.docx")
        
        assert result.success is False
        assert len(result.errors) > 0
        
        # Check for syntax errors
        syntax_errors = [e for e in result.errors if e.error_type == LintErrorType.SYNTAX_ERROR]
        assert len(syntax_errors) > 0
    
    @pytest.mark.asyncio
    async def test_nested_tags(self):
        """Test handling of properly nested tags."""
        content = """
        {% if user %}
            {% if user.profile %}
                {% for skill in user.profile.skills %}
                    {% if skill.certified %}
                        Certified: {{ skill.name }}
                    {% endif %}
                {% endfor %}
            {% endif %}
        {% endif %}
        """
        
        docx_bytes = self.create_test_docx(content)
        result = await self.linter.lint_docx_file(docx_bytes, "test.docx")
        
        assert result.success is True
        assert len(result.errors) == 0
    
    @pytest.mark.asyncio
    async def test_excessive_nesting(self):
        """Test detection of excessive nesting."""
        # Create deeply nested structure
        content = "Start\n"
        for i in range(15):  # Exceed the nesting limit
            content += f"{{% if condition{i} %}}\n"
        content += "Deep content\n"
        for i in range(14, -1, -1):
            content += f"{{% endif %}}\n"
        
        docx_bytes = self.create_test_docx(content)
        result = await self.linter.lint_docx_file(docx_bytes, "test.docx")
        
        # Should have nested error warnings
        nested_errors = [e for e in result.errors if e.error_type == LintErrorType.NESTED_ERROR]
        assert len(nested_errors) > 0
    
    @pytest.mark.asyncio
    async def test_long_line_warning(self):
        """Test detection of long lines."""
        # Create a very long line
        long_content = "This is a very long line that exceeds the maximum line length limit and should generate a warning about being too long for readability purposes and should be broken into smaller lines."
        
        options = LintOptions(max_line_length=50)
        docx_bytes = self.create_test_docx(long_content)
        result = await self.linter.lint_docx_file(docx_bytes, "test.docx", options)
        
        # Should have long line warnings
        long_line_warnings = [w for w in result.warnings if w.warning_type == LintWarningType.LONG_LINE]
        assert len(long_line_warnings) > 0
    
    @pytest.mark.asyncio
    async def test_empty_document(self):
        """Test handling of empty document."""
        docx_bytes = self.create_test_docx("")
        result = await self.linter.lint_docx_file(docx_bytes, "empty.docx")
        
        assert result.success is False
        assert len(result.errors) > 0
        
        # Should have document error for empty content
        doc_errors = [e for e in result.errors if e.error_type == LintErrorType.DOCUMENT_ERROR]
        assert len(doc_errors) > 0
    
    @pytest.mark.asyncio
    async def test_linting_options(self):
        """Test different linting options."""
        content = """
        This is a very long line that should trigger a warning when max_line_length is set low.
        {% if condition %}
        Content
        {% endif %}
        """
        
        # Test with warnings as errors
        options = LintOptions(
            max_line_length=30,
            fail_on_warnings=True
        )
        
        docx_bytes = self.create_test_docx(content)
        result = await self.linter.lint_docx_file(docx_bytes, "test.docx", options)
        
        # Should fail due to warnings being treated as errors
        if len(result.warnings) > 0:
            assert result.success is False
    
    @pytest.mark.asyncio
    async def test_completeness_score(self):
        """Test completeness score calculation."""
        # Good template
        good_content = """
        Hello {{ name }}!
        {% if company %}
        Company: {{ company }}
        {% endif %}
        """
        
        docx_bytes = self.create_test_docx(good_content)
        result = await self.linter.lint_docx_file(docx_bytes, "good.docx")
        
        assert result.summary.completeness_score is not None
        assert result.summary.completeness_score > 80  # Should be high for good template
        
        # Poor template with errors
        bad_content = """
        Hello {{ name }!  # Syntax error
        {% if condition %  # Syntax error
        """
        
        docx_bytes = self.create_test_docx(bad_content)
        result = await self.linter.lint_docx_file(docx_bytes, "bad.docx")
        
        assert result.summary.completeness_score is not None
        assert result.summary.completeness_score < 80  # Should be lower due to errors


class TestLinterAPI:
    """Test cases for the FastAPI linter endpoint."""
    
    def setup_method(self):
        """Set up test client."""
        self.client = TestClient(app)
    
    def create_test_file(self, content: str, filename: str = "test.docx") -> Dict[str, Any]:
        """
        Create a test file upload for the API.
        
        Args:
            content: Text content for the document
            filename: Filename to use
            
        Returns:
            Dictionary suitable for files parameter in requests
        """
        # Create a simple docx file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
            temp_path = temp_file.name
        
        try:
            doc = Document()
            for line in content.split('\n'):
                if line.strip():
                    doc.add_paragraph(line)
            doc.save(temp_path)
            
            # Read file content
            with open(temp_path, 'rb') as f:
                file_content = f.read()
            
            return {
                'document': (filename, file_content, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            }
        finally:
            if os.path.exists(temp_path):
                os.unlink(temp_path)
    
    def test_lint_valid_template(self):
        """Test API with valid template."""
        content = """
        Hello {{ name }}!
        {% if company %}
        Company: {{ company }}
        {% endif %}
        """
        
        files = self.create_test_file(content)
        
        response = self.client.post(
            "/api/v1/lint-docx-template",
            files=files
        )
        
        assert response.status_code == 200
        result = response.json()
        
        assert result['success'] is True
        assert result['summary']['total_errors'] == 0
        assert 'template_preview' in result
    
    def test_lint_with_options(self):
        """Test API with custom linting options."""
        content = "This is a very long line that should trigger warnings when max_line_length is set low"
        
        files = self.create_test_file(content)
        data = {
            'options': json.dumps({
                'max_line_length': 30,
                'verbose': True,
                'fail_on_warnings': False
            })
        }
        
        response = self.client.post(
            "/api/v1/lint-docx-template",
            files=files,
            data=data
        )
        
        assert response.status_code == 200
        result = response.json()
        
        # Should have warnings for long line
        assert len(result['warnings']) > 0
    
    def test_lint_invalid_file_type(self):
        """Test API with invalid file type."""
        response = self.client.post(
            "/api/v1/lint-docx-template",
            files={'document': ('test.txt', b'Hello world', 'text/plain')}
        )
        
        assert response.status_code == 400
        result = response.json()
        assert result['error_type'] == 'invalid_file_type'
    
    def test_lint_no_file(self):
        """Test API with no file provided."""
        response = self.client.post("/api/v1/lint-docx-template")
        
        assert response.status_code == 422  # FastAPI validation error
    
    def test_lint_empty_file(self):
        """Test API with empty file."""
        response = self.client.post(
            "/api/v1/lint-docx-template", 
            files={'document': ('empty.docx', b'', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
        )
        
        assert response.status_code == 400
        result = response.json()
        assert result['error_type'] == 'empty_file'
    
    def test_lint_template_with_errors(self):
        """Test API with template containing errors."""
        content = """
        Hello {{ name }!  # Missing closing brace
        {% if condition %}
        Content
        # Missing {% endif %}
        """
        
        files = self.create_test_file(content)
        
        response = self.client.post(
            "/api/v1/lint-docx-template",
            files=files
        )
        
        assert response.status_code == 200
        result = response.json()
        
        assert result['success'] is False
        assert result['summary']['total_errors'] > 0
        assert len(result['errors']) > 0
        
        # Check error structure
        first_error = result['errors'][0]
        assert 'error_type' in first_error
        assert 'message' in first_error
        assert 'suggestion' in first_error
    
    def test_api_response_structure(self):
        """Test that API response matches expected schema."""
        content = "Hello {{ name }}!"
        files = self.create_test_file(content)
        
        response = self.client.post(
            "/api/v1/lint-docx-template",
            files=files
        )
        
        assert response.status_code == 200
        result = response.json()
        
        # Check required fields
        required_fields = ['success', 'errors', 'warnings', 'summary', 'template_preview']
        for field in required_fields:
            assert field in result
        
        # Check summary structure
        summary = result['summary']
        summary_fields = ['total_errors', 'total_warnings', 'template_size', 'lines_count', 'jinja_tags_count']
        for field in summary_fields:
            assert field in summary
        
        # Check errors structure (if any)
        for error in result['errors']:
            assert 'error_type' in error
            assert 'message' in error
        
        # Check warnings structure (if any)
        for warning in result['warnings']:
            assert 'warning_type' in warning
            assert 'message' in warning


if __name__ == "__main__":
    # Run tests
    pytest.main([__file__, "-v"])