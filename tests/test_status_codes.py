#!/usr/bin/env python3
"""
Test script to verify that linter endpoints return 200 OK for both success and error cases.
"""

import asyncio
import os
import sys
from pathlib import Path
import requests
import json

# Add the project root to the Python path
sys.path.append(str(Path(__file__).parent.parent))


def create_broken_template_file():
    """Create a broken .docx template for testing."""
    from docx import Document
    
    # Create a truly broken template
    doc = Document()
    doc.add_paragraph('Customer: {{ customer.name }')  # Missing closing brace
    doc.add_paragraph('{% if customer.vip %}')
    doc.add_paragraph('VIP Benefits')
    doc.add_paragraph('{% for item in items %}')
    doc.add_paragraph('- {{ item.name }}')
    # Missing {% endfor %} and {% endif %}
    
    os.makedirs('temp', exist_ok=True)
    doc.save('temp/broken_for_status_test.docx')
    return 'temp/broken_for_status_test.docx'


def create_valid_template_file():
    """Create a valid .docx template for testing."""
    from docx import Document
    
    doc = Document()
    doc.add_paragraph('Customer: {{ customer.name }}')
    doc.add_paragraph('{% if customer.vip %}')
    doc.add_paragraph('VIP Benefits')
    doc.add_paragraph('{% endif %}')
    doc.add_paragraph('Items:')
    doc.add_paragraph('{% for item in items %}')
    doc.add_paragraph('- {{ item.name }}')
    doc.add_paragraph('{% endfor %}')
    
    os.makedirs('temp', exist_ok=True)
    doc.save('temp/valid_for_status_test.docx')
    return 'temp/valid_for_status_test.docx'


def test_standalone_linter_status_codes():
    """Test standalone linter endpoint status codes."""
    print("🧪 Testing Standalone Linter Endpoint Status Codes")
    print("-" * 50)
    
    base_url = "http://localhost:8000"
    
    # Test 1: Valid template - should return 200 with PDF
    print("\n1. Valid template test:")
    valid_file = create_valid_template_file()
    
    try:
        with open(valid_file, "rb") as f:
            response = requests.post(
                f"{base_url}/api/v1/lint-docx-template",
                files={"document": ("valid_template.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
            )
        
        print(f"  Status Code: {response.status_code}")
        print(f"  Content-Type: {response.headers.get('content-type')}")
        print(f"  Expected: 200 with PDF")
        
        if response.status_code == 200:
            if response.headers.get('content-type') == 'application/pdf':
                print("  ✅ PASS: Valid template returns 200 with PDF")
            else:
                print("  ❌ FAIL: Valid template should return PDF")
        else:
            print(f"  ❌ FAIL: Expected 200, got {response.status_code}")
            
    except Exception as e:
        print(f"  ❌ ERROR: {e}")
    
    # Test 2: Broken template - should return 200 with PDF error report
    print("\n2. Broken template test (PDF response):")
    broken_file = create_broken_template_file()
    
    try:
        with open(broken_file, "rb") as f:
            response = requests.post(
                f"{base_url}/api/v1/lint-docx-template",
                files={"document": ("broken_template.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
            )
        
        print(f"  Status Code: {response.status_code}")
        print(f"  Content-Type: {response.headers.get('content-type')}")
        print(f"  Expected: 200 with PDF error report")
        
        if response.status_code == 200:
            if response.headers.get('content-type') == 'application/pdf':
                print("  ✅ PASS: Broken template returns 200 with PDF error report")
                # Save the error report for inspection
                with open('temp/error_report_status_test.pdf', 'wb') as f:
                    f.write(response.content)
                print(f"  📄 Error report saved: temp/error_report_status_test.pdf")
            else:
                print("  ❌ FAIL: Broken template should return PDF error report")
        else:
            print(f"  ❌ FAIL: Expected 200, got {response.status_code}")
            print(f"  Response: {response.text[:200]}")
            
    except Exception as e:
        print(f"  ❌ ERROR: {e}")
    
    # Test 3: Broken template with JSON response option - should return 200 with JSON
    print("\n3. Broken template test (JSON response):")
    
    try:
        with open(broken_file, "rb") as f:
            data = {
                "options": json.dumps({
                    "response_format": "json",
                    "verbose": True
                })
            }
            response = requests.post(
                f"{base_url}/api/v1/lint-docx-template",
                files={"document": ("broken_template.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                data=data
            )
        
        print(f"  Status Code: {response.status_code}")
        print(f"  Content-Type: {response.headers.get('content-type')}")
        print(f"  Expected: 200 with JSON")
        
        if response.status_code == 200:
            if response.headers.get('content-type') == 'application/json':
                result = response.json()
                print("  ✅ PASS: Broken template returns 200 with JSON")
                print(f"  📊 Success: {result.get('success', 'N/A')}")
                print(f"  📊 Errors: {result.get('summary', {}).get('total_errors', 'N/A')}")
            else:
                print("  ❌ FAIL: Should return JSON when requested")
        else:
            print(f"  ❌ FAIL: Expected 200, got {response.status_code}")
            
    except Exception as e:
        print(f"  ❌ ERROR: {e}")


def test_integrated_endpoint_status_codes():
    """Test integrated endpoint status codes."""
    print("\n\n🧪 Testing Integrated Endpoint Status Codes")
    print("-" * 50)
    
    base_url = "http://localhost:8000"
    
    # Test 1: Valid template with data - should return 200 with PDF document
    print("\n1. Valid template with data test:")
    valid_file = create_valid_template_file()
    
    try:
        data = {
            "request_data": json.dumps({
                "template_data": {
                    "customer": {"name": "John Doe", "vip": True},
                    "items": [{"name": "Product A"}]
                }
            })
        }
        
        with open(valid_file, "rb") as f:
            response = requests.post(
                f"{base_url}/api/v1/process-template-document",
                files={"file": ("valid_template.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                data=data
            )
        
        print(f"  Status Code: {response.status_code}")
        print(f"  Content-Type: {response.headers.get('content-type')}")
        print(f"  Expected: 200 with PDF document")
        
        if response.status_code == 200:
            if response.headers.get('content-type') == 'application/pdf':
                print("  ✅ PASS: Valid template returns 200 with processed PDF")
            else:
                print("  ❌ FAIL: Valid template should return processed PDF")
        else:
            print(f"  ❌ FAIL: Expected 200, got {response.status_code}")
            print(f"  Response: {response.text[:200]}")
            
    except Exception as e:
        print(f"  ❌ ERROR: {e}")
    
    # Test 2: Broken template with data - should return 200 with PDF error report
    print("\n2. Broken template with data test (PDF error report):")
    broken_file = create_broken_template_file()
    
    try:
        data = {
            "request_data": json.dumps({
                "template_data": {
                    "customer": {"name": "John Doe", "vip": True},
                    "items": [{"name": "Product A"}]
                }
            })
        }
        
        with open(broken_file, "rb") as f:
            response = requests.post(
                f"{base_url}/api/v1/process-template-document",
                files={"file": ("broken_template.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                data=data
            )
        
        print(f"  Status Code: {response.status_code}")
        print(f"  Content-Type: {response.headers.get('content-type')}")
        print(f"  Expected: 200 with PDF error report")
        
        if response.status_code == 200:
            if response.headers.get('content-type') == 'application/pdf':
                print("  ✅ PASS: Broken template returns 200 with PDF error report")
                with open('temp/integrated_error_report.pdf', 'wb') as f:
                    f.write(response.content)
                print(f"  📄 Error report saved: temp/integrated_error_report.pdf")
            else:
                print("  ❌ FAIL: Broken template should return PDF error report")
                print(f"  Response: {response.text[:200]}")
        else:
            print(f"  ❌ FAIL: Expected 200, got {response.status_code}")
            print(f"  This was the old behavior (400) - should now be 200!")
            print(f"  Response: {response.text[:200]}")
            
    except Exception as e:
        print(f"  ❌ ERROR: {e}")
    
    # Test 3: Broken template with JSON response option - should return 200 with JSON
    print("\n3. Broken template with data test (JSON response):")
    
    try:
        data = {
            "request_data": json.dumps({
                "template_data": {
                    "customer": {"name": "John Doe", "vip": True},
                    "items": [{"name": "Product A"}]
                },
                "linter_options": {
                    "response_format": "json",
                    "verbose": True
                }
            })
        }
        
        with open(broken_file, "rb") as f:
            response = requests.post(
                f"{base_url}/api/v1/process-template-document",
                files={"file": ("broken_template.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                data=data
            )
        
        print(f"  Status Code: {response.status_code}")
        print(f"  Content-Type: {response.headers.get('content-type')}")
        print(f"  Expected: 200 with JSON")
        
        if response.status_code == 200:
            if response.headers.get('content-type') == 'application/json':
                result = response.json()
                print("  ✅ PASS: Broken template returns 200 with JSON error response")
                print(f"  📊 Status: {result.get('status', 'N/A')}")
                print(f"  📊 Errors: {result.get('linting_results', {}).get('summary', {}).get('total_errors', 'N/A')}")
            else:
                print("  ❌ FAIL: Should return JSON when requested")
        else:
            print(f"  ❌ FAIL: Expected 200, got {response.status_code}")
            
    except Exception as e:
        print(f"  ❌ ERROR: {e}")


def main():
    """Run all status code tests."""
    print("🔬 Linter Status Code Tests")
    print("=" * 60)
    print("This test verifies that linter endpoints return 200 OK")
    print("for both successful validation and error cases.")
    print("🚀 Make sure the service is running on http://localhost:8000")
    print()
    
    try:
        # Test standalone linter endpoint
        test_standalone_linter_status_codes()
        
        # Test integrated endpoint
        test_integrated_endpoint_status_codes()
        
        print(f"\n🏁 Status code tests completed!")
        print("=" * 60)
        print("\n✅ Expected behavior:")
        print("- All linting results should return 200 OK")
        print("- PDF error reports for validation failures (default)")
        print("- JSON error responses when explicitly requested")
        print("- No more 400 Bad Request for linting errors")
        
    except KeyboardInterrupt:
        print("\n⏹️  Test interrupted by user")
    except Exception as e:
        print(f"\n❌ Test failed: {e}")
        import traceback
        traceback.print_exc()
    finally:
        # Clean up test files
        try:
            import glob
            for file in glob.glob('temp/*_status_test.docx'):
                os.remove(file)
        except:
            pass


if __name__ == "__main__":
    main()